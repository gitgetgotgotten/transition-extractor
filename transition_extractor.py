import streamlit as st
import json
import re
from collections import defaultdict, Counter
from docx import Document
import zipfile
import io
from typing import List, Dict, Tuple, Optional

def extract_text_from_docx(uploaded_file) -> str:
    """Extract text from uploaded .docx file"""
    try:
        doc = Document(uploaded_file)
        full_text = []
        
        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text.strip())
        
        # Also extract from tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text.strip())
        
        return '\n'.join(full_text)
    except Exception as e:
        st.error(f"Error reading document: {str(e)}")
        return ""

def process_document(uploaded_file):
    """Updated process_document function with improved transition extraction"""
    try:
        # Read the document
        doc = Document(uploaded_file)
        filename = uploaded_file.name
        
        # Extract text
        full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        
        # Find ALL markers in the document (for multiple articles)
        marker = "À savoir également dans votre département"
        all_triplets = []
        all_transitions = []
        
        # Split text into potential articles
        marker_positions = []
        start_pos = 0
        while True:
            pos = full_text.find(marker, start_pos)
            if pos == -1:
                break
            marker_positions.append(pos)
            start_pos = pos + 1
        
        debug_info = {
            'text_length': len(full_text),
            'has_marker': len(marker_positions) > 0,
            'marker_count': len(marker_positions),
            'found_transitions': [],
            'articles_processed': 0,
            'transition_lines_found': 0
        }
        
        if not marker_positions:
            return [], [], filename, debug_info
        
        # Process each article section
        for i, marker_pos in enumerate(marker_positions):
            # Find the main paragraph after this marker
            main_paragraph_start = marker_pos + len(marker)
            
            # Find the transitions marker for this article
            transitions_marker_index = full_text.find("Transitions :", main_paragraph_start)
            if transitions_marker_index == -1:
                continue
            
            # Find the end of this article (next marker or end of transitions section)
            next_marker_pos = marker_positions[i + 1] if i + 1 < len(marker_positions) else len(full_text)
            
            # Make sure transitions marker is before next article
            if transitions_marker_index > next_marker_pos:
                continue
            
            # Extract the main paragraph (between marker and "Transitions:")
            main_paragraph = full_text[main_paragraph_start:transitions_marker_index].strip()
            
            # Extract transitions section
            transitions_start = transitions_marker_index + len("Transitions :")
            transitions_end = next_marker_pos
            
            # Look for article number pattern to end transitions section
            transitions_section = full_text[transitions_start:transitions_end]
            next_article_match = re.search(r'\n\s*\d+\s+du\s+\d+/\d+', transitions_section)
            if next_article_match:
                transitions_section = transitions_section[:next_article_match.start()].strip()
            
            # Extract individual transitions using the improved function
            article_transitions = extract_transitions_from_section(transitions_section)
            debug_info['found_transitions'].extend(article_transitions)
            
            # Process each transition ONLY within this article's main paragraph
            for transition in article_transitions:
                # Create variations of the transition
                transition_variations = create_transition_variations(transition)
                
                # Extract triplets for this transition ONLY from this article's main paragraph
                triplets = extract_context_around_transition(main_paragraph, transition, transition_variations)
                all_triplets.extend(triplets)
                
                # Update progress for user feedback
                if len(article_transitions) > 10:  # Only show progress for large batches
                    progress = (article_transitions.index(transition) + 1) / len(article_transitions)
                    st.progress(progress, text=f"Processing transition {article_transitions.index(transition) + 1}/{len(article_transitions)}")
            
            all_transitions.extend(article_transitions)
            debug_info['articles_processed'] += 1
        
        debug_info['transition_lines_found'] = len(all_transitions)
        return all_triplets, all_transitions, filename, debug_info
        
    except Exception as e:
        st.error(f"Error processing {uploaded_file.name}: {str(e)}")
        return [], [], uploaded_file.name, {'error': str(e)}


def extract_transitions_from_section(transitions_section: str) -> List[str]:
    """Extract clean transitions from the transitions section at the end of articles"""
    transitions = []
    
    for line in transitions_section.split('\n'):
        line = line.strip()
        if line and line != "Transitions :" and not re.match(r'^\d+\s+du\s+\d+/\d+', line):
            # Clean up common prefixes/suffixes
            line = re.sub(r'^[-•\d\.\s\:]+', '', line).strip()
            # Remove trailing punctuation and spaces
            line = re.sub(r'[,\s]+$', '', line).strip()
            if len(line) > 2:
                transitions.append(line)
    
    return transitions

def create_transition_variations(transition: str) -> List[str]:
    """Create variations of a transition to handle different formats and punctuation"""
    variations = []
    
    # Original transition
    variations.append(transition)
    
    # Basic case variations
    variations.append(transition.lower())
    variations.append(transition.capitalize())
    
    # Handle "que" vs "qu'" - FIXED VERSION
    if "que" in transition.lower():
        # Replace "que" at word boundary with "qu'"
        var_with_apostrophe = re.sub(r'\bque\b', "qu'", transition, flags=re.IGNORECASE)
        if var_with_apostrophe != transition:  # Only add if it's different
            variations.append(var_with_apostrophe)
            variations.append(var_with_apostrophe.lower())
            variations.append(var_with_apostrophe.capitalize())
    
    # Handle "qu'" vs "que" (reverse case)
    if "qu'" in transition.lower():
        var_without_apostrophe = re.sub(r"\bqu'", "que ", transition, flags=re.IGNORECASE)
        if var_without_apostrophe != transition:
            variations.append(var_without_apostrophe)
            variations.append(var_without_apostrophe.lower())
    
    # Handle punctuation variations
    base_transition = transition.rstrip('.,!?;:')
    if base_transition != transition:
        variations.append(base_transition)
        variations.append(base_transition.lower())
    
    # Add version with comma at the end
    if not transition.endswith(','):
        variations.append(transition + ',')
        variations.append((transition + ',').lower())
    
    # Add version with period at the end
    if not transition.endswith('.'):
        variations.append(transition + '.')
        variations.append((transition + '.').lower())
    
    # Remove duplicates while preserving order
    unique_variations = []
    for var in variations:
        if var and var not in unique_variations:
            unique_variations.append(var)
    
    return unique_variations



def find_sentence_boundaries(text: str) -> List[int]:
    """Find sentence boundaries in text, handling various edge cases"""
    boundaries = [0]  # Start of text
    
    # Improved sentence boundary detection
    sentence_endings = re.finditer(r'[.!?]+(?:\s+|$)', text)
    
    for match in sentence_endings:
        end_pos = match.end()
        # Skip abbreviations and numbers
        before_match = text[max(0, match.start()-10):match.start()]
        if not re.search(r'\b(?:M|Mme|Dr|St|etc|vs|cf|p|pp|vol|n°|art)\.$', before_match, re.IGNORECASE):
            boundaries.append(end_pos)
    
    # Also add paragraph boundaries as potential sentence boundaries
    paragraph_breaks = re.finditer(r'\n\s*\n', text)
    for match in paragraph_breaks:
        boundaries.append(match.end())
    
    boundaries.append(len(text))  # End of text
    return sorted(list(set(boundaries)))


def extract_context_around_transition(main_paragraph: str, transition: str, transition_variations: List[str]) -> List[Dict]:
    """Extract exactly one sentence before and after each transition occurrence - FOCUSED DEBUG"""
    triplets = []
    
    # Only debug the "Enfin" transition
    if "Enfin" in transition:
        print(f"\n=== DEBUGGING ENFIN TRANSITION ===")
        print(f"Looking for: '{transition}'")
        print(f"Main paragraph length: {len(main_paragraph)}")
        
        # Check if "enfin" exists in the text at all
        enfin_pos = main_paragraph.lower().find("enfin")
        if enfin_pos == -1:
            print("❌ 'enfin' NOT FOUND in main paragraph at all!")
            print(f"Last 200 chars of paragraph: ...{main_paragraph[-200:]}")
            return []
        else:
            print(f"✅ Found 'enfin' at position {enfin_pos}")
            # Show context around enfin
            start = max(0, enfin_pos - 50)
            end = min(len(main_paragraph), enfin_pos + 100)
            print(f"Context: ...{main_paragraph[start:end]}...")
        
        # Check each variation
        print(f"Testing {len(transition_variations)} variations:")
        for i, var in enumerate(transition_variations[:3]):  # Only show first 3
            found = main_paragraph.lower().find(var.lower())
            print(f"  {i+1}. '{var}' -> {'FOUND' if found != -1 else 'NOT FOUND'}")
    
    # Find all transition positions in the text
    transition_positions = []
    
    for variation in transition_variations:
        if not variation.strip():
            continue
            
        text_lower = main_paragraph.lower()
        var_lower = variation.lower().strip()
        
        start_pos = 0
        while True:
            pos = text_lower.find(var_lower, start_pos)
            if pos == -1:
                break
            
            actual_text = main_paragraph[pos:pos + len(variation)]
            transition_positions.append((pos, pos + len(variation), actual_text, transition))
            start_pos = pos + 1
    
    # Only log summary for "Enfin"
    if "Enfin" in transition:
        print(f"Total matches found: {len(transition_positions)}")
        if len(transition_positions) == 0:
            print("❌ NO MATCHES - This is the problem!")
            return []
    
    # Remove duplicates and sort by position
    unique_positions = []
    for pos_info in transition_positions:
        is_duplicate = False
        for existing in unique_positions:
            if abs(existing[0] - pos_info[0]) < 5:
                is_duplicate = True
                break
        if not is_duplicate:
            unique_positions.append(pos_info)
    
    unique_positions.sort(key=lambda x: x[0])
    
    # Process each transition occurrence
    for trans_start, trans_end, actual_transition, original_transition in unique_positions:
        # Find exactly one sentence before the transition
        text_before = main_paragraph[:trans_start]
        sentences_before = re.split(r'(?<=[.!?])\s+', text_before.strip())
        sentences_before = [s.strip() for s in sentences_before if s.strip()]
        
        if sentences_before:
            para_a_text = sentences_before[-1].strip()
        else:
            para_a_text = text_before.strip()
        
        if para_a_text and not para_a_text.endswith(('.', '!', '?')):
            para_a_text += '.'
        
        # Find exactly one sentence after the transition
        text_after = main_paragraph[trans_end:].strip()
        text_after = re.sub(r'^[,\s]+', '', text_after)
        
        sentence_match = re.search(r'^[^.!?]*[.!?](?=\s|$)', text_after)
        
        if sentence_match:
            para_b_text = sentence_match.group().strip()
        else:
            first_part = text_after.split('\n')[0]
            if len(first_part) > 100:
                para_b_text = first_part[:100].strip() + '.'
            else:
                para_b_text = first_part.strip()
                if para_b_text and not para_b_text.endswith(('.', '!', '?')):
                    para_b_text += '.'
        
        # Validate minimum content length
        if len(para_a_text) < 10 or len(para_b_text) < 10:
            continue
        
        # Create triplet
        triplet = {
            'paragraph_a': para_a_text,
            'transition': original_transition,
            'paragraph_b': para_b_text
        }
        
        # Only log result for "Enfin"
        if "Enfin" in transition:
            print(f"✅ CREATED ENFIN TRIPLET:")
            print(f"  A: '{triplet['paragraph_a'][:50]}...'")
            print(f"  B: '{triplet['paragraph_b'][:50]}...'")
        
        # Check for duplicates
        is_duplicate = False
        for existing in triplets:
            if (existing['paragraph_a'] == para_a_text and 
                existing['paragraph_b'] == para_b_text and
                existing['transition'] == original_transition):
                is_duplicate = True
                break
        
        if not is_duplicate:
            triplets.append(triplet)
    
    return triplets


def generate_outputs(all_triplets, all_transitions):
    """Generate various output formats from the extracted data."""
    # Count transition occurrences
    transition_counts = Counter([t['transition'] for t in all_triplets])
    
    # Cap each transition at 3 uses
    capped_triplets = []
    transition_usage = Counter()
    
    for triplet in all_triplets:
        transition = triplet['transition']
        if transition_usage[transition] < 3:
            capped_triplets.append(triplet)
            transition_usage[transition] += 1
    
    # 1. fewshot_examples.json
    fewshot_json = json.dumps(capped_triplets, indent=2, ensure_ascii=False)
    
    # 2. fewshots_rejected.txt
    rejected_transitions = []
    for transition, count in transition_counts.items():
        if count > 3:
            rejected_transitions.append(f"{transition}: {count}")
    fewshots_rejected_txt = "\n".join(rejected_transitions)
    
    # 3. transitions_only.txt
    transitions_txt = "\n".join(sorted(set(all_transitions)))
    
    # 4. transitions_only_rejected.txt
    transition_counts_all = Counter(all_transitions)
    transitions_rejected = []
    for transition, count in transition_counts_all.items():
        if count > 1:
            transitions_rejected.append(f"{transition}: {count}")
    transitions_only_rejected_txt = "\n".join(transitions_rejected)
    
    # 5. fewshot_examples.jsonl
    jsonl_examples = []
    for triplet in capped_triplets:
        example = {
            "messages": [
                {
                    "role": "system",
                    "content": "You are a helpful assistant that continues text based on the given context."
                },
                {
                    "role": "user",
                    "content": f"{triplet['paragraph_a']} {triplet['transition']}"
                },
                {
                    "role": "assistant",
                    "content": triplet['paragraph_b']
                }
            ]
        }
        jsonl_examples.append(json.dumps(example, ensure_ascii=False))
    
    fewshot_jsonl = "\n".join(jsonl_examples)
    
    # 6. fewshots-fineTuning_rejected.txt
    finetuning_rejected = []
    for transition, count in transition_counts.items():
        if count > 3:
            finetuning_rejected.append(f"{transition}: {count}")
    fewshots_finetuning_rejected_txt = "\n".join(finetuning_rejected)
    
    return (
        fewshot_json, 
        transitions_txt, 
        fewshot_jsonl, 
        fewshots_rejected_txt,
        transitions_only_rejected_txt,
        fewshots_finetuning_rejected_txt,
        len(capped_triplets)
    )

def create_download_zip(fewshot_json, transitions_txt, fewshot_jsonl, 
                       fewshots_rejected_txt, transitions_only_rejected_txt, 
                       fewshots_finetuning_rejected_txt):
    """Create a ZIP file containing all output files."""
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, 'w') as zip_file:
        zip_file.writestr('fewshot_examples.json', fewshot_json)
        zip_file.writestr('transitions_only.txt', transitions_txt)
        zip_file.writestr('fewshot_examples.jsonl', fewshot_jsonl)
        zip_file.writestr('fewshots_rejected.txt', fewshots_rejected_txt)
        zip_file.writestr('transitions_only_rejected.txt', transitions_only_rejected_txt)
        zip_file.writestr('fewshots-fineTuning_rejected.txt', fewshots_finetuning_rejected_txt)
    
    buffer.seek(0)
    return buffer.getvalue()

def main():
    # Initialize session state variables if they don't exist
    if 'all_triplets' not in st.session_state:
        st.session_state['all_triplets'] = []
    
    if 'all_transitions' not in st.session_state:
        st.session_state['all_transitions'] = []
    
    if 'processed_files' not in st.session_state:
        st.session_state['processed_files'] = []
    
    if 'debug_info' not in st.session_state:
        st.session_state['debug_info'] = []
    
    if 'outputs' not in st.session_state:
        st.session_state['outputs'] = {}
    
    st.set_page_config(
        page_title="Transition Extractor",
        page_icon="📄",
        layout="wide"
    )
    
    st.title("📄 Transition Extractor for News Articles")
    st.markdown("Extract structured transition examples from .docx news articles")
    
    # File upload
    st.header("1. Upload Documents")
    uploaded_files = st.file_uploader(
        "Choose .docx files",
        type=['docx'],
        accept_multiple_files=True,
        help="Upload one or more .docx news articles containing the marker 'À savoir également dans votre département'"
    )
    
    if uploaded_files:
        st.success(f"Uploaded {len(uploaded_files)} file(s)")
        
        # Process documents
        if st.button("🔍 Process Documents", type="primary"):
            with st.spinner("Processing documents..."):
                all_triplets = []
                all_transitions = []
                processed_files = []
                debug_info_all = []
                
                progress_bar = st.progress(0)
                
                for i, uploaded_file in enumerate(uploaded_files):
                    triplets, transitions, filename, debug_info = process_document(uploaded_file)
                    all_triplets.extend(triplets)
                    all_transitions.extend(transitions)
                    processed_files.append({
                        'filename': filename,
                        'triplets_count': len(triplets),
                        'transitions_count': len(transitions)
                    })
                    debug_info_all.append({
                        'filename': filename,
                        **debug_info
                    })
                    
                    progress_bar.progress((i + 1) / len(uploaded_files))
                
                # Store results in session state
                st.session_state['all_triplets'] = all_triplets
                st.session_state['all_transitions'] = all_transitions
                st.session_state['processed_files'] = processed_files
                st.session_state['debug_info'] = debug_info_all
                
                st.success("✅ Processing complete!")
    
    # Show results if available
    if 'all_triplets' in st.session_state and st.session_state['all_triplets']:
        st.header("2. Processing Results")
        
        # Display summary
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.metric("Total Triplets Found", len(st.session_state['all_triplets']))
        
        with col2:
            st.metric("Unique Transitions", len(set(st.session_state['all_transitions'])))
        
        with col3:
            st.metric("Files Processed", len(st.session_state['processed_files']))
        
        # Show per-file results
        st.subheader("Per-File Results")
        for file_info in st.session_state['processed_files']:
            st.write(f"**{file_info['filename']}**: {file_info['triplets_count']} triplets, {file_info['transitions_count']} transitions")
        
        # In the debug information section, update to show new fields:
        if st.session_state.get('debug_info'):
            with st.expander("🔍 Debug Information (Click to expand)"):
                for debug in st.session_state['debug_info']:
                    st.write(f"**{debug['filename']}**:")
                    st.write(f"- Text length: {debug['text_length']} characters")
                    st.write(f"- Markers found: {debug.get('marker_count', 0)}")
                    st.write(f"- Articles processed: {debug.get('articles_processed', 0)}")
                    st.write(f"- Total transitions found: {debug['transition_lines_found']}")
                    
                    # Show found transitions
                    if debug['found_transitions']:
                        st.write("**Found transitions:**")
                        for i, trans in enumerate(debug['found_transitions'][:20], 1):  # Show first 20
                            st.write(f"{i}. {trans}")
                        if len(debug['found_transitions']) > 20:
                            st.write(f"... and {len(debug['found_transitions']) - 20} more")
                    
                    # Show raw text preview only if no articles were processed
                    if debug.get('articles_processed', 0) == 0 and debug['text_length'] > 0:
                        st.text_area(
                            f"Raw text preview for {debug['filename']}:",
                            debug.get('raw_text_preview', '')[:500] + "..." if len(debug.get('raw_text_preview', '')) > 500 else debug.get('raw_text_preview', ''),
                            height=200,
                            key=f"debug_{debug['filename']}"
                        )
                    st.write("---")
        
        # Generate outputs
        st.header("3. Generate Outputs")
        
        output_formats = st.multiselect(
            "Select output formats to generate:",
            [
                'fewshot_examples.json', 
                'transitions_only.txt', 
                'fewshot_examples.jsonl', 
                'fewshots_rejected.txt',
                'transitions_only_rejected.txt',
                'fewshots-fineTuning_rejected.txt'
            ],
            default=[
                'fewshot_examples.json', 
                'transitions_only.txt', 
                'fewshot_examples.jsonl', 
                'fewshots_rejected.txt',
                'transitions_only_rejected.txt',
                'fewshots-fineTuning_rejected.txt'
            ]
        )
        
        if st.button("🔄 Generate Outputs"):
            with st.spinner("Generating outputs..."):
                fewshot_json, transitions_txt, fewshot_jsonl, fewshots_rejected_txt, \
                transitions_only_rejected_txt, fewshots_finetuning_rejected_txt, valid_examples = generate_outputs(
                    st.session_state['all_triplets'], 
                    st.session_state['all_transitions']
                )
                
                st.session_state['outputs'] = {
                    'fewshot_json': fewshot_json,
                    'transitions_txt': transitions_txt,
                    'fewshot_jsonl': fewshot_jsonl,
                    'fewshots_rejected_txt': fewshots_rejected_txt,
                    'transitions_only_rejected_txt': transitions_only_rejected_txt,
                    'fewshots_finetuning_rejected_txt': fewshots_finetuning_rejected_txt,
                    'valid_examples': valid_examples
                }
                
                st.success(f"✅ Generated outputs with {valid_examples} valid examples!")
    
    # Download section
    if 'outputs' in st.session_state and st.session_state['outputs']:
        st.header("4. Download Results")
        
        outputs = st.session_state['outputs']
        
        # Individual downloads
        col1, col2 = st.columns(2)
        
        with col1:
            if 'fewshot_examples.json' in output_formats:
                st.download_button(
                    "📄 Download fewshot_examples.json",
                    outputs['fewshot_json'],
                    "fewshot_examples.json",
                    "application/json"
                )
            
            if 'transitions_only.txt' in output_formats:
                st.download_button(
                    "📄 Download transitions_only.txt",
                    outputs['transitions_txt'],
                    "transitions_only.txt",
                    "text/plain"
                )
            
            if 'fewshot_examples.jsonl' in output_formats:
                st.download_button(
                    "📄 Download fewshot_examples.jsonl",
                    outputs['fewshot_jsonl'],
                    "fewshot_examples.jsonl",
                    "application/jsonl"
                )
        
        with col2:
            if 'fewshots_rejected.txt' in output_formats:
                st.download_button(
                    "📄 Download fewshots_rejected.txt",
                    outputs['fewshots_rejected_txt'],
                    "fewshots_rejected.txt",
                    "text/plain"
                )
            
            if 'transitions_only_rejected.txt' in output_formats:
                st.download_button(
                    "📄 Download transitions_only_rejected.txt",
                    outputs['transitions_only_rejected_txt'],
                    "transitions_only_rejected.txt",
                    "text/plain"
                )
            
            if 'fewshots-fineTuning_rejected.txt' in output_formats:
                st.download_button(
                    "📄 Download fewshots-fineTuning_rejected.txt",
                    outputs['fewshots_finetuning_rejected_txt'],
                    "fewshots-fineTuning_rejected.txt",
                    "text/plain"
                )
        
        # ZIP download
        st.subheader("Download All Files")
        zip_data = create_download_zip(
            outputs['fewshot_json'],
            outputs['transitions_txt'], 
            outputs['fewshot_jsonl'],
            outputs['fewshots_rejected_txt'],
            outputs['transitions_only_rejected_txt'],
            outputs['fewshots_finetuning_rejected_txt']
        )
        
        st.download_button(
            "📦 Download All Files (ZIP)",
            zip_data,
            "transition_extraction_results.zip",
            "application/zip"
        )
        
        # Preview section
        st.header("5. Preview Results")
        
        preview_format = st.selectbox(
            "Select format to preview:",
            [
                'fewshot_examples.json', 
                'transitions_only.txt', 
                'fewshot_examples.jsonl', 
                'fewshots_rejected.txt',
                'transitions_only_rejected.txt',
                'fewshots-fineTuning_rejected.txt'
            ]
        )
        
        if preview_format == 'fewshot_examples.json':
            st.code(outputs['fewshot_json'][:2000] + "..." if len(outputs['fewshot_json']) > 2000 else outputs['fewshot_json'], language='json')
        elif preview_format == 'transitions_only.txt':
            st.text(outputs['transitions_txt'][:2000] + "..." if len(outputs['transitions_txt']) > 2000 else outputs['transitions_txt'])
        elif preview_format == 'fewshot_examples.jsonl':
            st.code(outputs['fewshot_jsonl'][:2000] + "..." if len(outputs['fewshot_jsonl']) > 2000 else outputs['fewshot_jsonl'], language='json')
        elif preview_format == 'fewshots_rejected.txt':
            st.text(outputs['fewshots_rejected_txt'][:2000] + "..." if len(outputs['fewshots_rejected_txt']) > 2000 else outputs['fewshots_rejected_txt'])
        elif preview_format == 'transitions_only_rejected.txt':
            st.text(outputs['transitions_only_rejected_txt'][:2000] + "..." if len(outputs['transitions_only_rejected_txt']) > 2000 else outputs['transitions_only_rejected_txt'])
        elif preview_format == 'fewshots-fineTuning_rejected.txt':
            st.text(outputs['fewshots_finetuning_rejected_txt'][:2000] + "..." if len(outputs['fewshots_finetuning_rejected_txt']) > 2000 else outputs['fewshots_finetuning_rejected_txt'])
        
        # Show sample triplets
        st.header("6. Sample Triplets")
        if st.session_state['all_triplets']:
            sample_size = min(5, len(st.session_state['all_triplets']))
            st.write(f"Showing {sample_size} sample triplets:")
            
            for i, triplet in enumerate(st.session_state['all_triplets'][:sample_size], 1):
                with st.expander(f"Triplet {i}: {triplet['transition']}"):
                    st.write("**Paragraph A:**")
                    st.write(f"'{triplet['paragraph_a']}'")
                    st.write("**Transition:**")
                    st.write(f"*{triplet['transition']}*")
                    st.write("**Paragraph B:**")
                    st.write(f"'{triplet['paragraph_b']}'")

if __name__ == "__main__":
    main()